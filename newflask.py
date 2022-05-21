from flask import Flask
from flask import request
from resume_parser import resumeparse
import numpy as np
import docx
from docx import Document
skillsData = np.load("finalDataTags.npy")
import docx2txt
import pdfplumber
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import string
import spacy
nlp = spacy.load("en_core_web_sm")


app = Flask(__name__)

@app.route('/')
def hello():
    return '<h1>Hello, World!</h1>'

@app.route('/cv-form', methods=['POST'])
async def form_example():
    if request.method == 'POST':
        #job_description is already a string, to be compared with the CV
        job_description = request.form.get('job_description')
        
        #CV file from the request
        user_cv = request.files['file']
        extension = user_cv.filename.split('.')

        #if type is PDF
        if extension[-1] == 'pdf':
            #change from docx format to text (string) format
            resume = readFromPDF(user_cv)
        else: 
            #if type is doc/docx
            resume = docx2txt.process(user_cv)

        #clean both the resume and the job description
        cleaned_resume = cleanString(resume)
        cleaned_job_description = cleanString(job_description)

        #save them both as docx files
        buildDocxFile(cleaned_resume, "resume.docx")
        buildDocxFile(cleaned_job_description, 'job_description.docx')

        #parse and math the resume with the job description
        data = parseAndMatchResume("resume.docx","job_description.docx", "profile.pdf")
        print(data)
        return data

def readFromPDF(file): 
    text = ''
    with pdfplumber.open(file) as pdf:
        for pdf_page in pdf.pages:
            text += pdf_page.extract_text()
    return text

def buildDocxFile(data, filePath):
    document = Document()
    p = document.add_paragraph(data)
    document.save(filePath)

def parseAndMatchResume(resumePath, jobPath, fileName):
    rawResultsResume, rawResultsJob = getResults(resumePath, jobPath)
    rawResultsResume["scores"] = calculatePoints(rawResultsResume, rawResultsJob, fileName)
    rawResultsResume["matchScore"] = calculateMatch(rawResultsResume["scores"])
    finalResults = finalArrangeData(rawResultsResume, fileName)
    return finalResults

def getResults(resume_path, job_description_path):
    resumeData = resumeparse.read_file(resume_path)
    resumeData["skills"] = parseSkills(resumeData["skills"])
    #print("Resume Skills:", resumeData["skills"])

    jobData = resumeparse.read_file(job_description_path)
    jobData["skills"] = parseSkills(jobData["skills"])
    #print("Job Skills:", jobData["skills"])

    skillsRes = finalizeSkillsDisplay(resumeData["skills"], jobData["skills"])
    resumeData["skillsData"] = skillsRes

    miscResults = finalizeMetaResults(resume_path, job_description_path, resumeData)

    resumeData["meta"] = miscResults

    return resumeData, jobData

def calculatePoints(resumeData, jobData, fileName):
    skills = checkSkillsScore(resumeData["skills"], jobData["skills"])
    data = {
        "skills": skills,
        "ats": calcAtsScore(resumeData, fileName, skills),
        "rfindings": checkRecruiterScore(resumeData)
    }
    return data

def checkSkillsScore(resumeSkills, JobSkills):
    resumeSkills = [skill.lower() for skill in resumeSkills]
    JobSkills = [skill.lower() for skill in JobSkills]
    eScore = 0
    for jSkill in JobSkills:
        if jSkill in resumeSkills:
            eScore += 1
            
    return {"exists": eScore, "not-exists" : len(JobSkills) - eScore, "total" : len(JobSkills)}

def calculateMatch(score):
    skillsMatchScore = (score["skills"]["exists"] / score["skills"]["total"]) * 0.5
    atsMatchScore = (score["ats"]["exists"] / score["ats"]["total"]) * 0.25
    rfindingsScore = (score["rfindings"]["exists"] / score["rfindings"]["total"]) * 0.25
    totalScore = skillsMatchScore + atsMatchScore + rfindingsScore
    return round(totalScore * 100)

def finalArrangeData(resumeData, fileName):
    data = {}
    data["totalScore"] = resumeData["matchScore"]
    data["scores"] = resumeData["scores"]
    data["ats"] = {
        "resumeSkillsMissing" : resumeData["scores"]["skills"]["not-exists"],
        "educationMatch" : {"required": resumeData["meta"]["educationRequirements"], "match": False},
        "headings" : {"educationHeading" : resumeData["meta"]["educationHeading"], "workExperienceHeading" : resumeData["meta"]["workExperienceHeading"]},
        "dateFormatting" : True
    }
    if fileName != None:
        if fileName.split(".")[-1] in ["pdf", "docx"]:
            noSpecialCharName = True
            for i in ["-", "@", "!", "$", "^", "&", "*"]:
                if i in fileName:
                    noSpecialCharName = False
            data["ats"]["fileFormat"] = {
                "format" : fileName.split(".")[-1],
                "noSpecialCharName" : noSpecialCharName,
                "nameReadable" : True
            }
    data["recruiterFindings"] = {
        "length" : {"current" : resumeData["meta"]["length"], "allowed" : 1000},
        "measureableResults" : None,
        "wordsToAvoid" : None,
        "jobLevelMatch" : True
    }
    data["skillsData"] = resumeData["skillsData"]

    return data

def calcAtsScore(resumeData, fileName, skills):
    points = {}
    points["skills"] = 1 if skills["not-exists"] == 0 else 0

    if ("experience" in resumeData) and (resumeData["experience"] != None):
        points["experience"] = 1 if len(resumeData["experience"]) > 0 else 0
    else:
        points["experience"] = 0

    if ("email" in resumeData) and (resumeData["email"] != None):
        points["email"] = 1
    else:
        points["email"] = 0

    if ("mobile_number" in resumeData) and (resumeData["mobile_number"] != None):
        points["mobile_number"] = 1
    else:
        points["mobile_number"] = 0

    if ("name" in resumeData) and (resumeData["name"] != None):
        points["name"] = 1
    else:
        points["name"] = 0
        
    if resumeData["meta"]["educationRequirements"]:
        if resumeData["meta"]["educationHeading"]:
            points["educationReq"] = 1
    else:
        points["educationReq"] = 1
    points["educationHeading"] = 1 if resumeData["meta"]["educationHeading"] else 0
    points["workHeading"] = 1 if resumeData["meta"]["workExperienceHeading"] else 0
    if fileName != None:
        if fileName.split(".")[-1] in ["pdf", "docx"]:
            points["filePoint"] = 1
            points["fileNamePoint"] = 1
            for i in ["-", "@", "!", "$", "^", "&", "*"]:
                if i in fileName:
                    points["fileNamePoint"] = 0
            if points["fileNamePoint"] == 1:
                points["readAble"] = 1

    totalScore = 0
    totalItems = len(points.keys())
    for val in points.values():
        totalScore += val
    p = {
        "exists": totalScore,
        "not-exists": totalItems - totalScore,
        "total": totalItems
    }
    return p

def checkRecruiterScore(resumeData):
    points = {}
    points["wordCount"] = 1 if resumeData["meta"]["length"] < 1000 else 0
    points["measureableResults"] = 1
    points["avoidWords"] = 0
    points["jobLevel"] = 1
    i, total = 0, 0
    for val in points.values():
        total += val
    p = {
        "exists": total,
        "not-exists": 4 - total,
        "total": 4
    }
    return p


def parseSkills(skills):
    skillsFinal = []
    for skill in skills:
        lowerSkill = skill.lower()
        if lowerSkill in skillsData:
            skillsFinal.append(skill)
    return skillsFinal

def finalizeSkillsDisplay(resumeSkills, JobSkills):
    fSkills = {}
    resumeSkills = [skill.lower() for skill in resumeSkills]

    JobSkills = [skill.lower() for skill in JobSkills]
    eSkills = 0
    for skill in resumeSkills:
        if skill not in JobSkills:
            fSkills[skill] = {
                "wanted": False,
                "exists": True
            }
        else:
            eSkills += 1
            fSkills[skill] = {
                "wanted": True,
                "exists": True
            }
    for skill in JobSkills:
        if skill not in resumeSkills:
            fSkills[skill] = {
                "wanted": True,
                "exists": False
            }
        else:
            fSkills[skill] = {
                "wanted": True,
                "exists": False
            }
    return fSkills

def finalizeMetaResults(ResumePath, JobPath, ResumeData):
    meta = {}
    text = getTextFromDocx(ResumePath)
    meta["length"] = len(text)
    textJob = getTextFromDocx(JobPath)

    meta["educationHeading"] = True if checkEducation(text) else False    
    meta["workExperienceHeading"] = True if checkExperience(text) else False
    meta["educationRequirements"] = True if checkEducation(textJob) else False
    ResumeData["skillsData"] = countWords(ResumeData["skillsData"], text.lower(), textJob.lower())
    
    return meta

def getTextFromDocx(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def checkEducation(txt):
    edu = ["Education", "Degree"]
    for e in edu:
        if e.lower() in txt.lower():
            return True
    return False

def checkExperience(txt):
    exp = ["Experience", "Work Experience"]
    for e in exp:
        if e.lower() in txt.lower():
            return True
    return False

def checkEducation(txt):
    edu = ["Education", "Degree"]
    for e in edu:
        if e.lower() in txt.lower():
            return True
    return False

def countWords(skills, ResumeText, JobText):
    meta = {}
    for skill, val in skills.items():
        skillTest = skill.lower()
        meta[skill] = {
            "ResumeCount" : ResumeText.count(skillTest),
            "JobCount" : JobText.count(skillTest),
            "wanted" : val["wanted"],
            "exists" : val["exists"]
        }
    return meta

def cleanString(stringer): 
    cleaned_string = ''.join(c for c in stringer if valid_xml_char_ordinal(c))

    tokens = word_tokenize(stringer)

    # convert to lower case
    tokens = [w.lower() for w in tokens]

    # remove punctuation from each word
    table = str.maketrans(string.punctuation, ' ' * len(string.punctuation))
    stripped = [w.translate(table) for w in tokens]

    # remove remaining tokens that are not alphabetic
    words = [word for word in stripped if word.isalpha()]

    # filter out stop words    
    stop_words = set(stopwords.words('english'))
    words = [w for w in words if not w in stop_words]
    words = ' '.join(words)

    return str(words)

def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )